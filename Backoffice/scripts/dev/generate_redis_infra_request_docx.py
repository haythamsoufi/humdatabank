#!/usr/bin/env python3
"""Generate redis-infrastructure-request.docx for infra team."""

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Inches, Pt, RGBColor

GITHUB_BLOB_BASE = "https://github.com/IFRC-C2/IFRCNetworkDatabank/blob/main"

REFERENCE_LINKS: list[tuple[str, str]] = [
    (
        "redis-provisioning.md",
        f"{GITHUB_BLOB_BASE}/Backoffice/docs/runbooks/deployment/redis-provisioning.md",
    ),
    (
        "multi-instance-without-redis.md",
        f"{GITHUB_BLOB_BASE}/Backoffice/docs/runbooks/deployment/multi-instance-without-redis.md",
    ),
    (
        "azure-app-service.md (section 8: Multi-Worker)",
        f"{GITHUB_BLOB_BASE}/Backoffice/docs/runbooks/deployment/azure-app-service.md#8-multi-worker-considerations",
    ),
    (
        "gateway-loadtest-reproduction-plan.md",
        f"{GITHUB_BLOB_BASE}/Backoffice/docs/runbooks/incidents/gateway-loadtest-reproduction-plan.md",
    ),
    (
        "gateway-504-worker-saturation.md",
        f"{GITHUB_BLOB_BASE}/Backoffice/docs/runbooks/incidents/gateway-504-worker-saturation.md",
    ),
]


def _set_normal_style(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)


def _add_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)


def _add_subtitle(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x59, 0x59, 0x59)


def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def _add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def _add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def _add_hyperlink(paragraph, text: str, url: str) -> None:
    """Insert an external hyperlink run into an existing paragraph."""
    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    r_pr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    new_run.append(r_pr)
    text_elem = OxmlElement("w:t")
    text_elem.text = text
    new_run.append(text_elem)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def _add_reference_links(doc: Document, links: list[tuple[str, str]]) -> None:
    for label, url in links:
        p = doc.add_paragraph(style="List Bullet")
        _add_hyperlink(p, label, url)


def _add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    doc.add_paragraph()


def build_document() -> Document:
    doc = Document()
    _set_normal_style(doc)

    today = date.today().strftime("%d %B %Y")
    _add_title(doc, "Azure Managed Redis — Infrastructure Request")
    _add_subtitle(doc, "IFRC Network Databank Backoffice")
    _add_subtitle(doc, f"Prepared: {today}  |  Audience: Infrastructure / Platform team")

    doc.add_paragraph()

    _add_heading(doc, "1. Purpose", 1)
    doc.add_paragraph(
        "This document requests provision of Azure Managed Redis for the IFRC Network Databank "
        "Backoffice (App Service: ifrc-databank-app). It explains why Redis is needed for "
        "horizontal scalability, what it is used for in our application, what it is not used for, "
        "and how we propose to roll it out with minimal cost and risk."
    )

    _add_heading(doc, "2. Executive summary", 1)
    _add_bullets(
        doc,
        [
            "We are not asking for Redis to store user sessions. Login uses signed HTTP cookies; "
            "form data and business state live in PostgreSQL.",
            "We need Redis so that multiple App Service workers and instances can share small "
            "amounts of coordination state (rate limits, presence, alert cooldowns, diagnostics).",
            "Redis is a required part of our horizontal scale-out design: when App Service "
            "auto-scale adds or removes instances, Redis is how those instances coordinate and "
            "behave as one logical application instead of isolated silos.",
            "The Backoffice is not a basic single-server application. It is built as a "
            "multi-worker, multi-instance platform with Redis-backed coordination already "
            "implemented in code — running without Redis means operating in degraded fallback "
            "mode, not as architected.",
            "Without Redis we must keep ARR Affinity (sticky sessions) enabled and accept weakened "
            "security controls and degraded multi-user features when scaled out.",
            "Ongoing infra maintenance after go-live is almost nothing: Azure manages patching "
            "and HA; no backups, gateway changes, or manual Redis administration (see §10.3).",
            "Recommended Azure SKU: **Azure Managed Redis, Balanced B0 (1 GB), West Europe** — "
            "current Microsoft product for new workloads (see §3).",
            "Estimated cost (West Europe, CHF, public PAYG list, July 2026): ~CHF 11/month "
            "staging (single-node B0), ~CHF 22/month production (two-node HA). "
            "Confirm in Azure Pricing Calculator for your agreement.",
            "Compare: one App Service P2 v3 instance in West Europe is ~CHF 210–400/month — "
            "Redis is a small fraction of that while enabling safe scale-out.",
            "Load and traffic are increasing, not plateauing. Planned next phases include "
            "migrating the Indicator Bank and FDRS reporting systems into the Network Databank "
            "platform — Redis should be in place before that growth lands on a multi-instance "
            "deployment without shared coordination.",
        ],
    )

    _add_heading(doc, "3. Azure Managed Redis — product and SKU", 1)
    doc.add_paragraph(
        "Request provision of Azure Managed Redis (not the legacy Azure Cache for Redis "
        "product, which retires 30 September 2028). Pricing page: "
        "azure.microsoft.com/pricing/details/managed-redis/"
    )
    _add_table(
        doc,
        ["Setting", "Staging", "Production"],
        [
            ["Product", "Azure Managed Redis", "Azure Managed Redis"],
            ["Tier / SKU", "Balanced B0 (1 GB)", "Balanced B0 (1 GB)"],
            ["Region", "West Europe", "West Europe"],
            ["High availability", "Single-node (dev/test)", "Two-node HA (recommended)"],
            ["Networking", "Private Link to App Service VNet", "Private Link to App Service VNet"],
            ["App connection", "REDIS_URL=rediss://<host>:6380/0", "Same"],
        ],
    )
    _add_bullets(
        doc,
        [
            "The application connects via standard REDIS_URL (redis-py) — no Azure-specific SDK.",
            "Our Redis usage is small coordination keys (rate limits, presence ZSETs, alert "
            "cooldowns) — kilobytes to low megabytes, not session storage.",
            "B0 (1 GB) is more than sufficient; we do not need B1+, Premium, Modules, "
            "persistence, or clustering.",
            "Connection count is low (~3 Gunicorn workers × N instances).",
        ],
    )
    doc.add_paragraph(
        "Legacy note: Azure Cache for Redis (Standard C0, 250 MB) also works via REDIS_URL, "
        "but we are not requesting it for new provisioning because that product is retiring."
    )

    _add_heading(doc, "4. Response to infra feedback", 1)

    _add_heading(doc, '4.1 “Redis is only for systems with millions of sessions”', 2)
    doc.add_paragraph(
        "This statement reflects a common misconception. Redis is often associated with "
        "session storage at very large scale, but that is only one use case — and it is not "
        "why we need it for the IFRC Network Databank Backoffice."
    )
    doc.add_paragraph(
        "Redis is not only useful for systems with millions of sessions. It solves performance "
        "and coordination problems at almost any scale. Even a small application can use Redis "
        "for caching expensive database queries, storing sessions, rate limiting, temporary data, "
        "background job queues, distributed locks, and real-time features."
    )
    doc.add_paragraph(
        "The decision to adopt Redis should be based on latency requirements, traffic patterns, "
        "data lifetime, and operational complexity — not an arbitrary session count. At lower "
        "scale, Redis may improve response times and reduce database load. At very small scale, "
        "though, it can also be unnecessary infrastructure; the real question is whether the "
        "benefits justify running and maintaining it."
    )
    doc.add_paragraph("How this applies to our Backoffice specifically:")
    _add_bullets(
        doc,
        [
            "User authentication does not use server-side session storage. Flask-Login uses "
            "signed cookie sessions that work across all workers and instances without Redis.",
            "All durable application data (forms, submissions, users, notifications, AI documents) "
            "is stored in PostgreSQL.",
            "We are not requesting Redis because of session volume (tens to low hundreds of "
            "concurrent users, not millions). We need it because we run multiple Gunicorn workers "
            "× multiple App Service instances and those processes must share coordination state.",
            "Redis would hold only lightweight, TTL-bound coordination keys: rate-limit counters, "
            "presence membership, alert cooldown timestamps, optional cross-worker diagnostic snapshots.",
            "Typical working set is kilobytes to low megabytes — a shared notice board between "
            "workers, not a warehouse for session payloads.",
            "For our scale and architecture, the benefits (fleet-wide rate limits, correct "
            "multi-instance behaviour, safe auto-scale) justify the ~CHF 22/month production cost "
            "relative to the operational cost of running degraded fallback mode indefinitely.",
        ],
    )

    _add_heading(doc, '4.2 “Redis is expensive”', 2)
    doc.add_paragraph(
        "We agree cost must be justified. For our use case, Redis is one of the lowest-cost "
        "components on the platform relative to the alternatives."
    )

    _add_table(
        doc,
        ["Option", "West Europe — ~CHF / month (PAYG list)", "What it buys"],
        [
            [
                "Managed Redis Balanced B0 — staging (one node)",
                "~CHF 11",
                "Shared coordination; dev/test (no HA SLA)",
            ],
            [
                "Managed Redis Balanced B0 — production (two-node HA)",
                "~CHF 22",
                "Fleet-wide rate limits, presence, alert cooldown, diagnostics",
            ],
            [
                "App Service P2 v3 plan (one instance, comparison)",
                "~CHF 210–400",
                "More compute; does not fix per-worker coordination",
            ],
            [
                "No Redis + ARR Affinity On",
                "CHF 0 incremental",
                "Sticky routing, weakened limits, degraded co-editing presence",
            ],
        ],
    )

    doc.add_paragraph(
        "Source: Azure retail prices API, currencyCode=CHF, armRegionName=westeurope, "
        "July 2026. B0 meter CHF 0.0154/hour × 730 h ≈ CHF 11.24/month per node. "
        "Confirm final figures in the Azure Pricing Calculator for your agreement."
    )
    doc.add_paragraph(
        "Cost comparison: production Redis (~CHF 22/month) is roughly 5–10× cheaper than "
        "a single additional P2 v3 App Service instance, yet unlocks correct multi-instance behaviour."
    )
    _add_bullets(
        doc,
        [
            "We do not need B1, B3, Premium, Enterprise, or large Memory Optimized tiers.",
            "Reserved pricing (1- or 3-year) can reduce cost further once validated.",
        ],
    )

    _add_heading(doc, "5. What Redis is used for in our application", 1)
    doc.add_paragraph(
        "The application already supports Redis via a single environment variable: REDIS_URL. "
        "When set, the following subsystems use a shared Redis backend. When unset, each "
        "Gunicorn worker falls back to in-process memory (fail-open — the app keeps serving "
        "requests)."
    )

    _add_table(
        doc,
        ["Subsystem", "Without Redis", "With Redis", "User / ops impact"],
        [
            [
                "API & AI rate limiting",
                "Per worker (effective limit ≈ configured × workers × instances)",
                "Fleet-wide consistent limits",
                "Security: abuse limits are 6× weaker today on 2×3 workers",
            ],
            [
                "Security alert email cooldown",
                "Per worker",
                "Shared cooldown",
                "During 502/504 incidents, up to 6 duplicate alert emails observed",
            ],
            [
                "Form co-editing presence",
                "Per worker in-memory",
                "Shared Redis ZSET per assignment",
                "Co-editors on different workers may not see each other",
            ],
            [
                "AI WebSocket rate limits / daily quotas",
                "Per worker",
                "Shared counters",
                "Quota bypass possible by hitting different workers",
            ],
            [
                "504 / worker-pressure diagnostics",
                "Partial (filesystem mirror per container only)",
                "Cross-worker snapshot in Redis",
                "Faster root-cause analysis during gateway timeouts",
            ],
            [
                "Flask login sessions",
                "Signed cookies — already cross-worker",
                "No change (not stored in Redis)",
                "Not a driver for this request",
            ],
            [
                "Form submissions / business data",
                "PostgreSQL",
                "PostgreSQL",
                "Not a driver for this request",
            ],
        ],
    )

    _add_heading(doc, "6. Auto scale-out: why Redis is required with multiple instances", 1)
    doc.add_paragraph(
        "App Service auto-scale increases or decreases the number of running instances based on "
        "CPU, memory, or schedule. That is the intended way to handle reporting peaks, load tests, "
        "and growth without permanently paying for idle capacity. Auto-scale only delivers real "
        "benefit when every instance participates in the same application logic — not when each "
        "new instance is an isolated copy with its own private counters and state."
    )

    _add_heading(doc, "6.1 What happens when auto-scale adds an instance (without Redis)", 2)
    _add_bullets(
        doc,
        [
            "Azure starts a new App Service VM and boots 3 Gunicorn workers (current production setting).",
            "Each worker has its own in-memory rate-limit counters — a user or attacker can "
            "effectively multiply their allowed request rate by the number of workers × instances.",
            "Co-editing presence is stored per worker. Users on the new instance do not appear "
            "in the presence bar for users on the original instance, even though they edit the same form.",
            "Security alert cooldowns are per worker. A gateway incident can trigger duplicate "
            "alert emails from every worker on every instance.",
            "AI WebSocket connections and daily quotas are pinned to the worker that accepted "
            "the connection; auto-scale does not redistribute that state.",
            "ARR Affinity (sticky sessions) is required to paper over some of this, but affinity "
            "fights auto-scale: new instances receive less traffic until users are redistributed, "
            "and heavy users remain pinned to one instance while others sit under-utilized.",
        ],
    )

    _add_heading(doc, "6.2 What Redis enables during auto-scale (designed behaviour)", 2)
    _add_bullets(
        doc,
        [
            "All instances read and write the same coordination keys in Redis — rate limits, "
            "presence sets, alert cooldowns, and diagnostic snapshots are fleet-wide.",
            "When auto-scale adds instance #3 or removes instance #1, application behaviour "
            "stays consistent. No manual reconfiguration or affinity tuning is needed.",
            "ARR Affinity can be turned Off, so Application Gateway and the App Service load "
            "balancer distribute requests evenly across all healthy instances.",
            "Auto-scale becomes a capacity lever (more compute under load) rather than a "
            "correctness risk (fragmented state across silos).",
            "The application team can run controlled load tests and enable scale-out rules "
            "knowing security limits and multi-user features remain enforceable.",
        ],
    )

    doc.add_paragraph(
        "In short: auto-scale without Redis scales compute but not coordination. Redis is the "
        "shared coordination bus that makes multiple instances act as one system."
    )

    _add_heading(doc, "7. This is not a basic application — it is built for distributed operation", 1)
    doc.add_paragraph(
        "The IFRC Network Databank Backoffice is a multi-capability platform, not a simple "
        "monolithic web app that could run correctly on a single process with no shared state. "
        "From the start of the multi-worker / scale-out design, the engineering team implemented "
        "Redis-backed coordination with in-memory fallback — the code path activated by REDIS_URL "
        "is the intended production configuration."
    )

    _add_table(
        doc,
        ["Platform capability", "Why single-process memory is insufficient at scale"],
        [
            [
                "Real-time form co-editing presence",
                "Multiple users on different workers/instances must see the same live editor list",
            ],
            [
                "AI chat WebSockets and daily cost quotas",
                "Long-lived connections and spend limits must be enforceable fleet-wide",
            ],
            [
                "Authenticated API and AI rate limiting",
                "Abuse prevention must not multiply with every new worker or auto-scaled instance",
            ],
            [
                "Security monitoring and alert cooldown",
                "One incident should produce one alert, not N emails per worker",
            ],
            [
                "Background scheduler and digest jobs",
                "Partially mitigated via PostgreSQL advisory locks; Redis remains the standard "
                "pattern for fleet-wide job coordination",
            ],
            [
                "Gateway timeout diagnostics",
                "Cross-worker visibility during 502/504 incidents requires shared snapshots",
            ],
            [
                "PostgreSQL + pgvector AI/RAG",
                "Persistent data layer is separate; Redis handles ephemeral coordination only",
            ],
        ],
    )

    doc.add_paragraph(
        "Operating today without REDIS_URL is explicitly supported as a fail-open fallback "
        "(the app continues to serve users), but it is not the architected steady state for "
        "production at scale. Provision Redis is completing the infrastructure picture the "
        "application was written for — similar to how PostgreSQL is required for data, Redis is "
        "required for multi-instance coordination."
    )

    _add_heading(doc, "8. Why this matters for scalability", 1)
    doc.add_paragraph(
        "Scalability for us means running multiple App Service instances and multiple Gunicorn "
        "workers per instance without correctness or security regressions."
    )
    _add_bullets(
        doc,
        [
            "Today, scaling out to 2 instances requires ARR Affinity = On so users stay on the "
            "same instance. Affinity reduces effective load balancing: heavy users stay pinned, "
            "light endpoints queue behind them, and gateway timeouts become more likely under stress.",
            "With Redis, ARR Affinity can be turned Off. Requests distribute evenly across "
            "workers and instances while shared counters and presence remain consistent.",
            "Rate limiting and AI cost quotas must be enforceable fleet-wide before we expose "
            "the platform to higher concurrent usage or additional instances.",
            "Redis does not replace other performance work (Gunicorn tuning, WebSocket thread "
            "budget, blocking endpoint fixes). It is complementary: it removes structural "
            "blockers to horizontal scale and auto-scale.",
        ],
    )

    _add_heading(doc, "8.1 Growth trajectory — Indicator Bank and FDRS migrations", 2)
    doc.add_paragraph(
        "The IFRC Network Databank Backoffice is an expanding platform, not a static workload. "
        "Usage, reporting cycles, and concurrent editing load have been growing; upcoming "
        "roadmap phases will consolidate additional systems into this application."
    )
    _add_bullets(
        doc,
        [
            "Indicator Bank migration — central indicator definitions, history, embeddings, and "
            "lookup traffic will increasingly run through this Backoffice rather than separate "
            "tooling, adding API volume, admin UI load, and PostgreSQL/pgvector activity.",
            "FDRS (Federation-wide Reporting System) migration — form templates, submissions, "
            "validation workflows, and export pipelines for FDRS reporting will move into Network "
            "Databank, increasing peak traffic during reporting periods and background processing.",
            "Both migrations increase the need for fleet-wide rate limits, shared presence during "
            "co-editing, and safe App Service auto-scale — exactly the coordination Redis provides.",
            "Provisioning Redis now (~CHF 22/month production) avoids retrofitting coordination "
            "under higher load and gives staging a realistic environment to validate migrations "
            "before they reach production traffic.",
        ],
    )
    doc.add_paragraph(
        "In short: we are not sizing Redis for today's user count alone. We are preparing "
        "infrastructure for a consolidating platform where load and traffic are only increasing."
    )

    _add_heading(doc, "9. What Redis does not solve", 1)
    doc.add_paragraph(
        "To set expectations, Redis is not proposed as a fix for Application Gateway 502/504 "
        "timeouts caused by Gunicorn worker recycle or thread saturation. Those are addressed "
        "by separate application-side workstreams (worker config, WebSocket budget, slow "
        "endpoint caching, boot reliability). Redis enables safe scale-out alongside that work."
    )

    _add_heading(doc, "10. Proposed infrastructure request", 1)

    _add_heading(doc, "10.1 Azure resources", 2)
    _add_table(
        doc,
        ["Resource", "Specification", "Notes"],
        [
            [
                "Staging Redis",
                "Azure Managed Redis Balanced B0 (1 GB), West Europe",
                "Single-node (no HA); ~CHF 11/month PAYG list",
            ],
            [
                "Production Redis",
                "Azure Managed Redis Balanced B0 (1 GB), West Europe",
                "Two-node HA; ~CHF 22/month PAYG list",
            ],
            [
                "Connectivity",
                "Private endpoint preferred; TLS in transit",
                "Provide connection string as REDIS_URL app setting (slot-sticky)",
            ],
            [
                "High availability",
                "Single-node acceptable for staging; discuss two-node HA for prod",
                "Workload is coordination metadata, not primary data store; fail-open if unavailable",
            ],
        ],
    )

    _add_heading(doc, "10.2 Application Gateway, App Service, and Redis (one-time setup)", 2)
    doc.add_paragraph(
        "In our architecture, App Service to Redis does not use Application Gateway. "
        "Microsoft documents separate paths for user traffic versus app-to-backend traffic."
    )
    _add_bullets(
        doc,
        [
            "Users to web app: Application Gateway (Layer 7 HTTP/HTTPS/WebSocket) to App Service "
            "(Microsoft Application Gateway FAQ).",
            "App Service to Redis: regional VNet integration plus Private Link/private endpoint "
            "(Microsoft App Service VNet integration; Azure Managed Redis Private Link docs).",
            "Microsoft tutorial Isolate back-end communication: outbound from App Service is "
            "routed into the VNet to reach private endpoints — not via Application Gateway.",
            "Azure sample web-app-redis-sql-db: App Service + VNet integration + private "
            "endpoints to Redis — no Application Gateway in the Redis path.",
        ],
    )
    doc.add_paragraph(
        "Note: Application Gateway v2 can expose optional TCP/TLS listeners for some inbound "
        "non-HTTP workloads, but that is not the Microsoft pattern for App Service calling "
        "Redis and is not part of our setup. We do not configure an AGW listener for Redis."
    )
    _add_table(
        doc,
        ["Component", "Action for Redis", "Ongoing infra work?"],
        [
            [
                "Application Gateway",
                "No configuration change required",
                "None — Redis is not on the gateway path",
            ],
            [
                "Application Gateway backend timeout (≥300s for AI/SSE)",
                "Unchanged — already needed for long AI streams; unrelated to Redis",
                "None (existing setting)",
            ],
            [
                "Private Link (App Service VNet ↔ Redis)",
                "Provision with Redis resource",
                "Only if VNet layout changes later (rare)",
            ],
            [
                "App Service REDIS_URL",
                "Set once per deployment slot",
                "Only if secret rotates (rare)",
            ],
            [
                "App Service ARR Affinity",
                "Set Off after Redis verified (was On without Redis)",
                "Leave Off — no further changes",
            ],
        ],
    )
    doc.add_paragraph(
        "Infra takeaway: provision Redis + Private Link, set REDIS_URL, flip ARR Affinity Off "
        "after validation. No Application Gateway rule updates, listeners, or backend pool "
        "changes are needed for Redis."
    )

    _add_heading(doc, "10.3 Ongoing maintenance — infrastructure after go-live", 2)
    doc.add_paragraph(
        "After the one-time setup above, infra has almost nothing to do. Azure Managed Redis "
        "is fully managed — not a second database to operate. There are no recurring backup "
        "jobs, patching runbooks, Redis CLI sessions, or Application Gateway changes."
    )
    _add_table(
        doc,
        ["Responsibility", "Who", "Routine effort"],
        [
            [
                "Redis engine updates, OS patching, HA failover",
                "Azure (managed service)",
                "Zero — automatic",
            ],
            [
                "Backups / restore",
                "Not required",
                "N/A — TTL coordination keys only; no business data",
            ],
            [
                "Application Gateway / WAF for Redis",
                "Not applicable",
                "Zero — Redis bypasses the gateway",
            ],
            [
                "Private Link, REDIS_URL",
                "Infrastructure",
                "Near zero — only on rare network or secret rotation",
            ],
            [
                "Optional Azure Monitor glance (memory, availability)",
                "Infrastructure",
                "Passive — same as other PaaS resources",
            ],
            [
                "App verification after releases",
                "Application team",
                "Occasional — not infra day-to-day",
            ],
            [
                "Manual Redis CLI, key curation, SKU scale-up",
                "Nobody / unlikely for years",
                "Not required in normal operations",
            ],
        ],
    )
    doc.add_paragraph(
        "If Redis is briefly unavailable (e.g. Azure planned maintenance), the application "
        "keeps serving users — fail-open to in-memory per worker. Login and PostgreSQL data "
        "are unaffected. Infra should treat Redis as provision-and-forget with optional "
        "monitoring, not an ongoing operational burden."
    )

    _add_heading(doc, "11. Rollout and validation plan", 1)
    _add_numbered(
        doc,
        [
            "Provision Redis for staging; connect REDIS_URL on the staging App Service slot.",
            "Place staging behind Application Gateway (same timeout class as production) if not already done.",
            "Run baseline load test without Redis changes (document 502/504 rate, latency).",
            "Enable Redis on staging; re-run the same load test (before/after comparison).",
            "Run fail-open test: stop Redis mid-test and confirm the application continues serving (degraded shared state, no outage).",
            "Review metrics: Redis memory use, connection count, error rate.",
            "Provision production Redis; enable REDIS_URL on production during a low-traffic window.",
            "Turn ARR Affinity Off on production after 48 hours of clean metrics.",
            "Enable App Service auto-scale rules on staging; confirm behaviour with 2+ instances and Redis connected.",
            "Share a short test report with infra (1–2 pages: config, results, recommendation).",
        ],
    )

    _add_heading(doc, "12. Risk if Redis is not provisioned", 1)
    _add_table(
        doc,
        ["Risk", "Severity", "Current mitigation without Redis"],
        [
            [
                "Weakened rate limiting under scale-out",
                "High",
                "Manual tuning; divide limits by worker count; keep instance count low",
            ],
            [
                "Duplicate security alert emails during incidents",
                "Medium",
                "Accept noise; triage via logs not email count",
            ],
            [
                "Inaccurate co-editing presence",
                "Low (UX)",
                "ARR Affinity On; users may not see co-editors on other workers",
            ],
            [
                "Auto-scale adds instances without shared coordination",
                "High",
                "Avoid auto-scale or accept fragmented rate limits, presence, and alerts",
            ],
            [
                "Cannot turn off ARR Affinity",
                "Medium",
                "Must keep sticky sessions indefinitely; undermines load distribution",
            ],
            [
                "Scheduler duplication",
                "Low (mitigated)",
                "PostgreSQL advisory locks deployed 2026-07-23",
            ],
        ],
    )

    _add_heading(doc, "13. Summary ask", 1)
    doc.add_paragraph(
        "Please provision Azure Managed Redis Balanced B0 (1 GB) in West Europe for staging "
        "and production, with Private Link to the App Service VNet. Staging: single-node; "
        "production: two-node high availability. Estimated ~CHF 11/month (staging) and "
        "~CHF 22/month (production) at public PAYG list prices."
    )
    doc.add_paragraph(
        "Redis is required so that App Service auto-scale and multi-instance deployment work "
        "as designed: instances coordinate through a shared store, security limits apply "
        "fleet-wide, and ARR Affinity can be disabled for proper load balancing. The "
        "application is already built for this architecture — provision Redis completes the "
        "platform stack (PostgreSQL for data, Redis for coordination)."
    )

    _add_heading(doc, "14. References (internal engineering docs)", 1)
    doc.add_paragraph(
        f"Repository: {GITHUB_BLOB_BASE.replace('/blob/main', '')}"
    )
    _add_reference_links(doc, REFERENCE_LINKS)

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("Contact: IFRC Network Databank application team")
    run.italic = True

    return doc


def main() -> None:
    repo_root = Path(__file__).resolve().parents[3]
    out_dir = repo_root / "Backoffice" / "docs" / "runbooks" / "deployment"
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / "redis-infrastructure-request.docx"
    tmp_path = out_dir / "redis-infrastructure-request-updated.docx"

    doc = build_document()
    try:
        doc.save(str(out_path))
        print(f"Wrote {out_path}")
    except PermissionError:
        doc.save(str(tmp_path))
        print(f"Wrote {tmp_path} (close open docx and rename, or re-run)")


if __name__ == "__main__":
    main()

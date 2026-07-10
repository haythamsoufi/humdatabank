"""Tests for Word table border styling in render_docx.py."""

from __future__ import annotations

import sys
import tempfile
import unittest
from pathlib import Path
from unittest.mock import patch

from docx import Document
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / "scripts"))

from pb_figures.render_docx import (  # noqa: E402
    _add_donut_pair_block,
    _set_table_inner_borders,
)


class TableInnerBorderTests(unittest.TestCase):
    def test_outer_borders_removed_and_inner_borders_grey(self) -> None:
        doc = Document()
        table = doc.add_table(rows=2, cols=2)
        _set_table_inner_borders(table)

        borders = table._tbl.tblPr.find(qn("w:tblBorders"))
        self.assertIsNotNone(borders)

        for edge in ("top", "left", "bottom", "right"):
            element = borders.find(qn(f"w:{edge}"))
            self.assertIsNotNone(element)
            self.assertEqual(element.get(qn("w:val")), "nil")

        for edge in ("insideH", "insideV"):
            element = borders.find(qn(f"w:{edge}"))
            self.assertIsNotNone(element)
            self.assertEqual(element.get(qn("w:val")), "single")
            self.assertEqual(element.get(qn("w:color")), "C0C0C0")
            self.assertEqual(element.get(qn("w:sz")), "4")


class DonutPairTableTests(unittest.TestCase):
    def test_donut_pair_uses_single_four_column_table(self) -> None:
        doc = Document()
        with tempfile.TemporaryDirectory() as tmp:
            assets_dir = Path(tmp)
            items = [
                {"label": "Left label", "value": 44, "value_label": "44"},
                {"label": "Right label", "value": 194, "value_label": "194"},
            ]

            def fake_render(item, output_path, **kwargs):
                output_path.write_bytes(b"fake")

            with patch("pb_figures.render_docx.render_donut_asset", side_effect=fake_render), patch(
                "pb_figures.render_docx._add_donut_image_cell",
            ):
                _add_donut_pair_block(doc, items, assets_dir, "SP4_pair")

        self.assertEqual(len(doc.tables), 1)
        table = doc.tables[0]
        self.assertEqual(len(table.rows), 1)
        self.assertEqual(len(table.rows[0].cells), 4)
        self.assertEqual(table.cell(0, 0).text, "Left label")
        self.assertEqual(table.cell(0, 2).text, "Right label")

        layout = table._tbl.tblPr.find(qn("w:tblLayout"))
        self.assertIsNotNone(layout)
        self.assertEqual(layout.get(qn("w:type")), "fixed")

        tr_pr = table.rows[0]._tr.trPr
        self.assertIsNotNone(tr_pr)
        self.assertIsNotNone(tr_pr.find(qn("w:cantSplit")))


if __name__ == "__main__":
    unittest.main()

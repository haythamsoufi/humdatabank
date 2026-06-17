"""
Generate Architecture Review Board report on Backoffice automated CI/CD testing.
Output: docs/ARB-Backoffice-Automated-Testing-Report.docx
"""
from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

REPO_ROOT = Path(__file__).resolve().parents[2]
OUTPUT = REPO_ROOT / "docs" / "ARB-Backoffice-Automated-Testing-Report.docx"


def _set_normal_style(doc: Document) -> None:
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)


def _add_title_page(doc: Document) -> None:
    for _ in range(4):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Architecture Review Board\nSubmission")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    doc.add_paragraph()

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Automated CI/CD Test Coverage\nHumanitarian Databank — Backoffice")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)

    doc.add_paragraph()
    doc.add_paragraph()

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"Report date: {date.today().strftime('%d %B %Y')}\n")
    meta.add_run("Classification: Internal — Architecture Review\n")
    meta.add_run("Recommendation: Approval")

    doc.add_page_break()


def _heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def _para(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold


def _bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def _table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for p in hdr_cells[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for r_idx, row in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = val


def build_report() -> Document:
    doc = Document()
    _set_normal_style(doc)
    _add_title_page(doc)

    # 1. Executive Summary
    _heading(doc, "1. Executive Summary")
    _para(
        doc,
        "The Humanitarian Databank Backoffice implements a mature, production-grade automated "
        "testing and CI/CD programme that provides comprehensive regression protection across "
        "core humanitarian data workflows. The programme combines a large-scale pytest suite, "
        "PostgreSQL-faithful integration testing, layered security scanning, and continuous "
        "quality gates integrated into every pull request affecting Backoffice code.",
    )
    _para(doc, "Key assurance highlights:", bold=True)
    _bullets(
        doc,
        [
            "13,883 automated test cases across 370+ test modules, executed in parallel on every qualifying pull request.",
            "Three parallel CI jobs: fast code guards, Bandit static security analysis, and full pytest with coverage reporting.",
            "PostgreSQL 14 service container in CI ensures test environment parity with production.",
            "Complementary repository controls: CodeQL (required merge gate), Gitleaks secret scanning, and dependency review.",
            "Deep coverage of forms, assignments, indicator bank, RBAC, APIs, mobile integration, exports, and data quality.",
        ],
    )
    _para(
        doc,
        "This submission demonstrates that automated testing adequately covers core Backoffice "
        "functions and aligns with enterprise architecture standards for reliability, security, "
        "and maintainability. Approval is recommended.",
    )

    # 2. Purpose and Scope
    _heading(doc, "2. Purpose and Scope")
    _para(
        doc,
        "This report documents the automated test strategy and CI/CD integration for the "
        "Backoffice component of the Humanitarian Databank monorepo, submitted for Architecture "
        "Review Board approval. The Backoffice is the authoritative system for form design, "
        "indicator stewardship, country assignments, multilingual data collection, administrative "
        "governance, and partner/mobile APIs.",
    )
    _heading(doc, "2.1 In Scope", level=2)
    _bullets(
        doc,
        [
            "Unified pytest test suite (Backoffice/tests/)",
            "GitHub Actions CI pipeline (backoffice-ci.yml)",
            "Security and quality gates applied to Backoffice changes",
            "Test infrastructure: fixtures, factories, markers, and coverage reporting",
        ],
    )
    _heading(doc, "2.2 Complementary Assurance", level=2)
    _bullets(
        doc,
        [
            "CodeQL static analysis across Python and JavaScript/TypeScript",
            "On-demand k6 and Azure load testing for staging performance validation",
            "Developer handbook and contributing guidelines documenting test conventions",
        ],
    )

    # 3. CI/CD Architecture
    _heading(doc, "3. CI/CD Pipeline Architecture")
    _para(
        doc,
        "The Backoffice CI workflow triggers automatically on pull requests to main and develop "
        "whenever Backoffice code changes. Three independent jobs run in parallel, maximising "
        "feedback speed while maintaining thorough validation.",
    )
    _table(
        doc,
        ["CI Job", "Purpose", "Outcome"],
        [
            [
                "Code guards",
                "Diff-based checks for unsafe Jinja/JS patterns (CSP, inline handlers, eval)",
                "Prevents security and maintainability regressions in templates",
            ],
            [
                "Bandit (Python security)",
                "Static analysis of Backoffice/app with configured severity thresholds",
                "Continuous SAST coverage on application code",
            ],
            [
                "pytest suite",
                "Full automated test run with PostgreSQL 14, parallel workers, and coverage XML",
                "Validates business logic, routes, services, and API contracts",
            ],
        ],
    )
    doc.add_paragraph()
    _para(doc, "CI test environment configuration:", bold=True)
    _bullets(
        doc,
        [
            "Python 3.11 with cached dependency installation",
            "FLASK_CONFIG=testing with dedicated TEST_DATABASE_URL",
            "Parallel execution via pytest-xdist (-n auto)",
            "JUnit XML and coverage XML artefacts retained for audit and trend analysis",
        ],
    )

    # 4. Test Framework
    _heading(doc, "4. Test Framework and Organisation")
    _para(
        doc,
        "The test suite follows industry best practices with clear layering, shared infrastructure, "
        "and domain-aligned organisation. Tests are discoverable, maintainable, and suitable for "
        "both local developer workflows and automated CI execution.",
    )
    _table(
        doc,
        ["Layer", "Modules", "Role"],
        [
            ["Unit", "~323", "Fast isolated tests for routes, services, models, middleware, and forms"],
            ["Integration", "29", "Database-backed flows: auth, entry form, admin, RBAC, exports"],
            ["API", "24+", "REST v1 and mobile API contract validation"],
        ],
    )
    doc.add_paragraph()
    _para(doc, "Test infrastructure capabilities:", bold=True)
    _bullets(
        doc,
        [
            "Central conftest.py with PostgreSQL schema lifecycle management and role-based fixtures",
            "Deterministic test factories (factories.py) with parallel-safe unique identifiers",
            "Pytest markers: unit, integration, api, critical, transaction, auth_security, and more",
            "Interactive test runner (run_tests.bat) and documented coverage workflows",
            "Documented entry-form coverage matrix (93+ scenarios in ENTRY_FORM_TEST_COVERAGE.md)",
        ],
    )
    _para(
        doc,
        "Coverage policy: minimum 50%, target 70%+, with 90%+ emphasis on critical paths. "
        "CI produces machine-readable coverage reports on every run.",
    )

    # 5. Core Function Coverage
    _heading(doc, "5. Coverage of Core Business Functions")
    _para(
        doc,
        "Automated tests are mapped directly to humanitarian operations—not merely to technical "
        "layers—ensuring that the functions operators and focal points rely on are continuously "
        "validated.",
    )

    domains = [
        (
            "5.1 Identity, Access, and Governance",
            [
                "Session and extended authentication flows (integration and route-level tests)",
                "RBAC management with entity permissions (50+ dedicated integration scenarios)",
                "Authorization service internals and scoped access controls",
                "API key lifecycle and management routes",
                "Critical admin and focal-point route smoke tests with parametrized auth matrices",
                "Site lock and operational access controls",
            ],
        ),
        (
            "5.2 Form Lifecycle (Design → Entry → Submit)",
            [
                "Form builder: sections, items, templates, cloning, KoBo import pathways",
                "Entry form: 77+ integration tests plus dedicated API and route suites",
                "FormDataService, FormProcessingService, and TemplatePreparationService coverage",
                "Dynamic indicators, repeat sections, variable resolution, and localization",
                "Public submissions and preview mode validation",
                "Excel and PDF export route verification",
            ],
        ),
        (
            "5.3 Assignments and Country Operations",
            [
                "Assignment CRUD, workflow, and NS review authorisation",
                "Assignment lifecycle and completion services",
                "Country and entity services with extended edge-case coverage",
                "Document carryover, upload, download, and permission enforcement",
            ],
        ),
        (
            "5.4 Indicator Bank and Reference Data",
            [
                "Indicator bank administration and tab navigation",
                "System admin: sectors, lookups, and countries",
                "FDRS and data-quality methodology validation",
                "Indicator resolution and bank service logic",
            ],
        ),
        (
            "5.5 APIs (Public, Partner, and Mobile)",
            [
                "REST v1: users, data, countries, assignments, submissions, documents",
                "AI chat endpoint contracts",
                "Mobile JWT authentication and device routes",
                "Mobile admin surfaces: users, org, content, analytics, requests, notifications",
                "CSRF protection on API surfaces",
            ],
        ),
        (
            "5.6 Platform Reliability and Security",
            [
                "Transaction middleware: commit/rollback semantics and failure handling",
                "Activity, session, and API tracking middleware (100+ activity scenarios)",
                "Security headers, rate limiting, and malicious file upload defences",
                "Email delivery, campaigns, and notification routes",
                "Plugin framework manager and form integration",
                "Error handlers, logging configuration, and startup tasks",
            ],
        ),
    ]

    for title, items in domains:
        _heading(doc, title, level=2)
        _bullets(doc, items)

    # 6. Security Posture
    _heading(doc, "6. Security and Quality Assurance Posture")
    _para(
        doc,
        "Security is addressed through multiple complementary automated layers, providing "
        "defence-in-depth appropriate for a system handling humanitarian indicator and form data.",
    )
    _table(
        doc,
        ["Control", "Automation", "Benefit"],
        [
            ["Bandit SAST", "Every Backoffice PR", "Detects common Python security anti-patterns"],
            ["Gitleaks", "Every PR", "Prevents credential leakage into the repository"],
            ["CodeQL", "Every PR; required on main", "Deep static analysis for Python and JS/TS"],
            ["Dependency Review", "Every PR", "Flags vulnerable dependency changes"],
            ["Inline JS / CSP guards", "Every Backoffice PR", "Blocks unsafe template patterns in diffs"],
            ["AuthZ regression tests", "Continuous", "RBAC, form auth, and API authentication suites"],
            ["Upload abuse tests", "Continuous", "Malicious file upload integration scenarios"],
        ],
    )

    # 7. Maturity and Strengths
    _heading(doc, "7. Programme Maturity and Architectural Strengths")
    _bullets(
        doc,
        [
            "Scale with discipline: 13,883 parametrized cases provide broad regression matrices without brittle full-browser dependency for server logic.",
            "Production parity: PostgreSQL-only testing eliminates environment-specific false confidence.",
            "Domain alignment: tests organised around assignments, forms, indicators, and governance—the actual operator workflows.",
            "Operational transparency: CI uploads JUnit and coverage artefacts for audit trails and trend monitoring.",
            "Developer enablement: documented markers, factories, parallel runs, and interactive test wizard lower the barrier to adding tests with new features.",
            "Monorepo integration: Backoffice CI coexists with CodeQL, security scanning, and load-test workflows for holistic assurance.",
            "Critical-path smoke tests: admin and focal dashboards validated on every integration run via dedicated critical markers.",
        ],
    )

    # 8. Governance Alignment
    _heading(doc, "8. Governance and Continuous Improvement")
    _para(
        doc,
        "The testing programme is governed by documented standards in Backoffice/tests/README.md, "
        "CONTRIBUTING.md, and the Developer Handbook. New features are expected to ship with "
        "accompanying tests; CI enforces this norm on every Backoffice pull request.",
    )
    _para(
        doc,
        "The architecture supports ongoing enhancement without compromising current assurance: "
        "marker-based fast subsets, append-mode local coverage, on-demand load testing, and "
        "artefact retention provide a solid foundation for evolving requirements as the "
        "humanitarian data platform grows.",
    )

    # 9. Conclusion
    _heading(doc, "9. Conclusion and Recommendation")
    _para(
        doc,
        "The Humanitarian Databank Backoffice automated CI/CD testing programme meets and exceeds "
        "expectations for an enterprise humanitarian data platform. It delivers:",
    )
    _bullets(
        doc,
        [
            "Comprehensive automated regression coverage across core business functions",
            "Layered security validation integrated into the development workflow",
            "Production-faithful integration testing on PostgreSQL",
            "Clear organisational structure, documentation, and maintainability patterns",
            "Continuous execution on every relevant code change with auditable artefacts",
        ],
    )
    _para(
        doc,
        "On the basis of the evidence presented in this submission, the Architecture Review Board "
        "is respectfully asked to approve the Backoffice automated testing and CI/CD architecture "
        "as fit for purpose, operationally mature, and aligned with organisational standards for "
        "reliable humanitarian data stewardship.",
        bold=True,
    )

    # Appendix
    doc.add_page_break()
    _heading(doc, "Appendix A — Key References")
    _table(
        doc,
        ["Artifact", "Location"],
        [
            ["Test guide", "Backoffice/tests/README.md"],
            ["Pytest configuration", "Backoffice/pytest.ini"],
            ["CI workflow", ".github/workflows/backoffice-ci.yml"],
            ["Entry form coverage", "Backoffice/tests/integration/ENTRY_FORM_TEST_COVERAGE.md"],
            ["Branch protection ruleset", ".github/ruleset-protect-main.json"],
            ["Contributing / CI overview", "CONTRIBUTING.md"],
            ["Developer handbook", "docs/DEVELOPER-HANDBOOK.md"],
        ],
    )
    doc.add_paragraph()
    _heading(doc, "Appendix B — Summary Metrics")
    _table(
        doc,
        ["Metric", "Value"],
        [
            ["Automated test cases (collected)", "13,883"],
            ["Test modules", "370+"],
            ["CI parallel jobs (Backoffice PR)", "3"],
            ["Database in CI", "PostgreSQL 14"],
            ["Coverage reporting", "XML + terminal on every CI run"],
            ["Security scans (complementary)", "Bandit, CodeQL, Gitleaks, Dependency Review"],
        ],
    )

    return doc


def main() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = build_report()
    doc.save(str(OUTPUT))
    print(f"Report written to: {OUTPUT}")


if __name__ == "__main__":
    main()

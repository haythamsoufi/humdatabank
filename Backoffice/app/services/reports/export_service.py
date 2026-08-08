"""Export report widget payloads to Excel, PDF, HTML, and Word."""

from __future__ import annotations

import base64
import io
import logging
import re
from typing import Any

from flask import render_template, url_for

from app.models import ReportDefinition, User
from app.services.reports.schema import migrate_v1_to_v2
from app.services.reports.translation_helpers import normalize_language, resolve_translation

logger = logging.getLogger(__name__)


class ReportExportService:
    @staticmethod
    def export_excel(report: ReportDefinition, run_result: dict[str, Any]) -> bytes:
        from openpyxl import Workbook
        from openpyxl.utils import get_column_letter

        wb = Workbook()
        summary = wb.active
        summary.title = "Summary"
        summary.append(["Report", report.title])
        summary.append(["Slug", report.slug])
        summary.append(["Language", run_result.get("language") or ""])
        summary.append([])

        widgets = run_result.get("widgets") or {}
        for idx, (wid, payload) in enumerate(widgets.items()):
            title = payload.get("title") or wid
            sheet_name = re.sub(r"[^A-Za-z0-9 _-]", "", title)[:28] or f"Widget{idx + 1}"
            if idx == 0 and summary.title == "Summary":
                ws = summary
            else:
                ws = wb.create_sheet(title=sheet_name)
            ws.append(["Widget", title])
            if payload.get("type") == "kpi":
                ws.append(["Value", payload.get("value")])
                continue
            if payload.get("type") in {"table"} or payload.get("rows"):
                columns = payload.get("columns") or []
                rows = payload.get("rows") or []
                if columns:
                    ws.append(list(columns))
                for row in rows:
                    if isinstance(row, dict):
                        ws.append([row.get(c) for c in columns] if columns else list(row.values()))
                    else:
                        ws.append(list(row))
                continue
            chart = payload.get("chart_payload") or {}
            if chart.get("type") == "line":
                ws.append(["Year", "Value"])
                for pt in chart.get("series") or []:
                    ws.append([pt.get("x"), pt.get("y")])
            elif chart.get("type") == "bar":
                ws.append(["Label", "Value"])
                for cat in chart.get("categories") or []:
                    ws.append([cat.get("label"), cat.get("value")])
            elif chart.get("type") == "pie":
                ws.append(["Label", "Value"])
                for sl in chart.get("slices") or []:
                    ws.append([sl.get("label"), sl.get("value")])

        for sheet in wb.worksheets:
            for col_idx, _ in enumerate(sheet.iter_cols(min_row=1, max_row=1), start=1):
                sheet.column_dimensions[get_column_letter(col_idx)].width = 18

        buf = io.BytesIO()
        wb.save(buf)
        return buf.getvalue()

    @staticmethod
    def export_html(
        report: ReportDefinition,
        run_result: dict[str, Any],
        *,
        chart_images: dict[str, str] | None = None,
        language: str = "en",
    ) -> bytes:
        html = render_template(
            "admin/reports/export_pdf.html",
            report=report,
            widgets=run_result.get("widgets") or {},
            chart_images=chart_images or {},
            language=language,
            rtl=normalize_language(language) == "ar",
        )
        return html.encode("utf-8")

    @staticmethod
    def export_pdf(
        report: ReportDefinition,
        run_result: dict[str, Any],
        *,
        chart_images: dict[str, str] | None = None,
        language: str = "en",
    ) -> bytes:
        chart_images = chart_images or {}
        html = render_template(
            "admin/reports/export_pdf.html",
            report=report,
            widgets=run_result.get("widgets") or {},
            chart_images=chart_images,
            language=language,
            rtl=normalize_language(language) == "ar",
        )
        try:
            from weasyprint import CSS, HTML  # type: ignore
        except Exception as exc:
            raise RuntimeError("PDF generation is not available on this deployment.") from exc

        rtl = normalize_language(language) == "ar"
        pdf_css = """
            @page { size: A4; margin: 18mm; }
            body { font-family: Arial, sans-serif; color: #111827; font-size: 11pt; }
            h1 { font-size: 18pt; margin-bottom: 8px; }
            h2 { font-size: 13pt; margin-top: 16px; border-bottom: 1px solid #ccc; }
            .widget { margin: 12px 0 20px; page-break-inside: avoid; }
            table { width: 100%; border-collapse: collapse; margin-top: 8px; }
            th, td { border: 1px solid #ddd; padding: 4px 6px; text-align: left; }
            th { background: #f3f4f6; }
            img { max-width: 100%; height: auto; }
        """
        if rtl:
            pdf_css += """
            body { direction: rtl; font-family: 'Tajawal', Arial, sans-serif; }
            th, td { text-align: right; }
            """
        return HTML(string=html).write_pdf(stylesheets=[CSS(string=pdf_css)])

    @staticmethod
    def export_docx(
        report: ReportDefinition,
        run_result: dict[str, Any],
        *,
        chart_images: dict[str, str] | None = None,
        language: str = "en",
    ) -> bytes:
        from docx import Document
        from docx.shared import Inches, Pt

        chart_images = chart_images or {}
        definition = migrate_v1_to_v2(report.definition_json or {})
        theme = definition.get("theme") or {}
        doc = Document()
        title = report.title
        doc.add_heading(title, level=0)
        if report.description:
            doc.add_paragraph(report.description)

        for section in run_result.get("sections") or []:
            section_title = section.get("title") or section.get("id")
            doc.add_heading(section_title, level=1)
            for widget in section.get("widgets") or []:
                wid = widget.get("id")
                payload = (run_result.get("widgets") or {}).get(wid) or {}
                wtitle = payload.get("title") or widget.get("title") or wid
                doc.add_heading(wtitle, level=2)
                if payload.get("type") == "kpi":
                    p = doc.add_paragraph()
                    p.add_run(str(payload.get("value", "—"))).bold = True
                elif payload.get("type") == "text":
                    doc.add_paragraph(payload.get("content") or "")
                elif payload.get("rows"):
                    columns = payload.get("columns") or list((payload.get("rows") or [{}])[0].keys())
                    table = doc.add_table(rows=1, cols=len(columns))
                    hdr = table.rows[0].cells
                    for idx, col in enumerate(columns):
                        hdr[idx].text = str(col)
                    for row in payload.get("rows") or []:
                        cells = table.add_row().cells
                        for idx, col in enumerate(columns):
                            cells[idx].text = str(row.get(col, "") if isinstance(row, dict) else "")
                img = chart_images.get(wid)
                if img and isinstance(img, str) and img.startswith("data:image"):
                    decoded = ReportExportService.decode_chart_images({wid: img}).get(wid)
                    if decoded:
                        stream = io.BytesIO(decoded)
                        doc.add_picture(stream, width=Inches(5.5))
                footnote = payload.get("footnote")
                if footnote:
                    fp = doc.add_paragraph(footnote)
                    fp.runs[0].font.size = Pt(9)

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    @staticmethod
    def decode_chart_images(raw: dict[str, Any] | None) -> dict[str, bytes]:
        if not raw:
            return {}
        decoded: dict[str, bytes] = {}
        for wid, data_url in raw.items():
            if not isinstance(data_url, str):
                continue
            match = re.match(r"^data:image/png;base64,(.+)$", data_url.strip())
            if not match:
                continue
            try:
                decoded[wid] = base64.b64decode(match.group(1))
            except Exception:
                continue
        return decoded

    @staticmethod
    def rasterize_charts_headless(report_id: int, language: str, user: User) -> dict[str, str]:
        """Best-effort server-side chart rasterization via Playwright when available."""
        try:
            from playwright.sync_api import sync_playwright
        except ImportError:
            logger.warning("Playwright not installed; publish will omit chart images for language %s", language)
            return {}

        from flask import current_app

        base_url = current_app.config.get("REPORT_BUILDER_PRINT_BASE_URL") or "http://127.0.0.1:5000"
        print_url = f"{base_url}/admin/reports/{report_id}/print?language={language}"
        images: dict[str, str] = {}
        try:
            with sync_playwright() as p:
                browser = p.chromium.launch(headless=True)
                page = browser.new_page(viewport={"width": 1280, "height": 900})
                page.goto(print_url, wait_until="networkidle", timeout=120_000)
                hosts = page.query_selector_all("[data-widget-id]")
                for host in hosts:
                    wid = host.get_attribute("data-widget-id")
                    if not wid:
                        continue
                    png = host.screenshot(type="png")
                    if png:
                        images[wid] = "data:image/png;base64," + base64.b64encode(png).decode("ascii")
                browser.close()
        except Exception:
            logger.exception("Headless chart rasterization failed for report %s language %s", report_id, language)
        return images

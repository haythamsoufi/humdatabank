"""Tests for Word table border styling in render_docx.py."""

from __future__ import annotations

import struct
import sys
import tempfile
import unittest
from pathlib import Path
from unittest.mock import patch

from docx import Document
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / "scripts"))

from pb_figures.line_chart import CHART_HEIGHT  # noqa: E402
from pb_figures.render_docx import (  # noqa: E402
    CHART_WIDTH_PX,
    _DOCX_CONTENT_WIDTH_IN,
    _DOCX_LABEL_COL_IN,
    _DOCX_PAGE_MARGIN,
    _add_cumulative_block,
    _add_donut_pair_block,
    _chart_area_width,
    _chart_render_width_px,
    _configure_page_margins,
    _cumulative_table_widths,
    _docx_chart_display_height_in,
    _docx_chart_font_scale,
    _docx_chart_height_px,
    _donut_pair_widths,
    _set_cell_text,
    _set_table_inner_borders,
)
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT


def _minimal_png_bytes(width: int = 4, height: int = 4) -> bytes:
    """Build the smallest PNG that python-docx's ``add_picture`` can size.

    ``docx.image.png`` only reads the IHDR chunk's width/height (and an
    optional pHYs chunk for DPI, defaulting to 72 when absent); it never
    decodes pixel data or validates chunk CRCs. So a signature + IHDR +
    IEND with a dummy CRC is enough, without a real encoder/decoder
    dependency in the test.
    """

    def chunk(tag: bytes, data: bytes) -> bytes:
        return struct.pack(">I", len(data)) + tag + data + b"\x00\x00\x00\x00"

    signature = b"\x89PNG\r\n\x1a\n"
    ihdr_data = struct.pack(">IIBBBBB", width, height, 8, 2, 0, 0, 0)
    return signature + chunk(b"IHDR", ihdr_data) + chunk(b"IEND", b"")


class TableInnerBorderTests(unittest.TestCase):
    def test_outer_borders_removed_and_inner_borders_grey(self) -> None:
        doc = Document()
        table = doc.add_table(rows=2, cols=2)
        _set_table_inner_borders(table)

        borders = table._tbl.tblPr.find(qn("w:tblBorders"))
        self.assertIsNotNone(borders)

        for edge in ("top", "left", "bottom", "right"):
            element = borders.find(qn(f"w:{edge}"))
            self.assertIsNotNone(element)
            self.assertEqual(element.get(qn("w:val")), "nil")

        for edge in ("insideH", "insideV"):
            element = borders.find(qn(f"w:{edge}"))
            self.assertIsNotNone(element)
            self.assertEqual(element.get(qn("w:val")), "single")
            self.assertEqual(element.get(qn("w:color")), "C0C0C0")
            self.assertEqual(element.get(qn("w:sz")), "4")

    def test_configure_page_margins_uses_narrow_margins(self) -> None:
        doc = Document()
        _configure_page_margins(doc)
        section = doc.sections[0]
        self.assertEqual(section.left_margin, _DOCX_PAGE_MARGIN)
        self.assertEqual(section.right_margin, _DOCX_PAGE_MARGIN)

    def test_set_cell_text_centers_vertically(self) -> None:
        doc = Document()
        table = doc.add_table(rows=1, cols=1)
        _set_cell_text(table.cell(0, 0), "Indicator label")
        cell = table.cell(0, 0)
        self.assertEqual(cell.vertical_alignment, WD_CELL_VERTICAL_ALIGNMENT.CENTER)
        self.assertEqual(cell.text, "Indicator label")

    def test_cumulative_table_widths_widen_indicator_column(self) -> None:
        widths = _cumulative_table_widths(5)
        self.assertEqual(widths[0], _DOCX_LABEL_COL_IN)
        self.assertGreater(widths[0], 2.05)

    def test_cumulative_table_widths_span_content_area(self) -> None:
        widths = _cumulative_table_widths(5)
        self.assertAlmostEqual(sum(widths), _DOCX_CONTENT_WIDTH_IN, places=5)
        self.assertAlmostEqual(_chart_area_width(5), _DOCX_CONTENT_WIDTH_IN - _DOCX_LABEL_COL_IN, places=5)
        self.assertGreater(widths[1], 0.9)

    def test_donut_pair_widths_span_content_area(self) -> None:
        widths = _donut_pair_widths()
        self.assertAlmostEqual(sum(widths), _DOCX_CONTENT_WIDTH_IN, places=5)

    def test_docx_chart_height_matches_png_aspect_ratio(self) -> None:
        width_px = _chart_render_width_px(5)
        height_px = _docx_chart_height_px(width_px)
        png_ratio = CHART_HEIGHT / CHART_WIDTH_PX
        self.assertAlmostEqual(height_px / width_px, png_ratio, places=2)
        width_in = _chart_area_width(5)
        height_in = _docx_chart_display_height_in(width_in)
        self.assertAlmostEqual(height_in / width_in, png_ratio, places=4)
        self.assertAlmostEqual(_docx_chart_font_scale(width_px), width_px / CHART_WIDTH_PX, places=4)
        self.assertGreater(_docx_chart_font_scale(width_px), 1.0)


class DonutPairTableTests(unittest.TestCase):
    def test_donut_pair_uses_single_four_column_table(self) -> None:
        doc = Document()
        with tempfile.TemporaryDirectory() as tmp:
            assets_dir = Path(tmp)
            items = [
                {"label": "Left label", "value": 44, "value_label": "44"},
                {"label": "Right label", "value": 194, "value_label": "194"},
            ]

            def fake_render(item, output_path, **kwargs):
                output_path.write_bytes(b"fake")

            with patch("pb_figures.render_docx.render_donut_asset", side_effect=fake_render), patch(
                "pb_figures.render_docx._add_donut_image_cell",
            ):
                _add_donut_pair_block(doc, items, assets_dir, "SP4_pair")

        self.assertEqual(len(doc.tables), 1)
        table = doc.tables[0]
        self.assertEqual(len(table.rows), 1)
        self.assertEqual(len(table.rows[0].cells), 4)
        self.assertEqual(table.cell(0, 0).text, "Left label")
        self.assertEqual(table.cell(0, 2).text, "Right label")

        layout = table._tbl.tblPr.find(qn("w:tblLayout"))
        self.assertIsNotNone(layout)
        self.assertEqual(layout.get(qn("w:type")), "fixed")

        tr_pr = table.rows[0]._tr.trPr
        self.assertIsNotNone(tr_pr)
        self.assertIsNotNone(tr_pr.find(qn("w:cantSplit")))


class CumulativeBlockRegressionTests(unittest.TestCase):
    """Regression test for a real production bug (fixed in this codebase's history):

    _add_cumulative_block used to reference `n_years` in
    `_chart_render_width_px(n_years)` several lines before `n_years` was
    assigned, raising UnboundLocalError for every cumulative indicator. It
    shipped to main and broke every DOCX build until caught by this kind of
    test — CI only runs render_docx's *unit-level* helpers (table widths,
    borders, etc.), never `_add_cumulative_block` itself, so the bug wasn't
    caught until a real build hit it.
    """

    def test_add_cumulative_block_available_item_does_not_raise(self) -> None:
        doc = Document()
        item = {
            "label": "Indicator label",
            "years": ["2021", "2022", "2023"],
            "reporting": ["10", "20", "30"],
            "implementing": ["5", "15", "25"],
        }
        labels = {"year": "Year", "reporting": "Reporting", "implementing": "Implementing"}

        def fake_render(item, target_label, output_path, **kwargs):
            output_path.write_bytes(_minimal_png_bytes())

        with tempfile.TemporaryDirectory() as tmp:
            assets_dir = Path(tmp)
            with patch("pb_figures.render_docx.render_line_chart_asset", side_effect=fake_render):
                _add_cumulative_block(doc, item, labels, "Target label", assets_dir, "SP1")

        self.assertEqual(len(doc.tables), 1)
        table = doc.tables[0]
        self.assertEqual(len(table.columns), len(item["years"]) + 1)
        # Year row (bold), reporting row, and implementing row are all present.
        self.assertEqual(len(table.rows), 4)
        self.assertEqual(table.cell(0, 0).text, item["label"])

    def test_add_cumulative_block_respects_ns_table_mode(self) -> None:
        """Same bug path, hit via the ns_table_mode branch that hides the implementing row."""
        doc = Document()
        item = {
            "label": "Indicator label",
            "years": ["2022", "2023"],
            "reporting": ["1", "2"],
            "implementing": ["1", "2"],
            "ns_table_mode": "ns_unit",
        }
        labels = {"year": "Year", "reporting": "Reporting", "implementing": "Implementing"}

        def fake_render(item, target_label, output_path, **kwargs):
            output_path.write_bytes(_minimal_png_bytes())

        with tempfile.TemporaryDirectory() as tmp:
            assets_dir = Path(tmp)
            with patch("pb_figures.render_docx.render_line_chart_asset", side_effect=fake_render):
                _add_cumulative_block(doc, item, labels, "Target label", assets_dir, "SP2")

        table = doc.tables[0]
        # Chart row + year row + reporting row only (no implementing row).
        self.assertEqual(len(table.rows), 3)


if __name__ == "__main__":
    unittest.main()

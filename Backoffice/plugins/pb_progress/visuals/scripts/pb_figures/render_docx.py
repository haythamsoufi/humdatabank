"""Render chart-only PNG assets and editable Word reports."""

from __future__ import annotations

import tempfile
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from .languages import ARABIC_VISUAL_FONT, LATIN_DOCX_FONT, is_rtl
from .calculations import not_available
from .layouts import cumulative_table_rows, mapping_from_model, section_has_indicators
from .line_chart import render_line_chart_svg
from .donut_chart import render_donut_svg
from .svg_raster import write_svg_png
from .payload import build_payload
from .report_meta import report_parts, report_titles, section_uses_part_heading_only

CHART_WIDTH_PX = 481
IFRC_RED = RGBColor(0xC2, 0x25, 0x26)
_DOCX_STYLES = ("Normal", "Title", "Heading 1", "Heading 2", "List Paragraph")
_DOCX_PAGE_MARGIN = Inches(0.5)
_DOCX_PAGE_WIDTH_IN = 8.5
_DOCX_CONTENT_WIDTH_IN = _DOCX_PAGE_WIDTH_IN - (2 * 0.5)
_DOCX_LABEL_COL_IN = 2.85
_DOCX_DONUT_COL_IN = 0.55
_DOCX_TABLE_FONT = 9
_DOCX_CHART_DPI = 175  # PNG resolution for embedded line charts (~481px @ 2.75in)


def _set_rfonts(r_pr, font_name: str) -> None:
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        r_fonts.set(qn(attr), font_name)


def _docx_font(language: str) -> str:
    return ARABIC_VISUAL_FONT if is_rtl(language) else LATIN_DOCX_FONT


def _apply_run_font(run, language: str) -> None:
    font_name = _docx_font(language)
    run.font.name = font_name
    _set_rfonts(run._element.get_or_add_rPr(), font_name)


def _apply_paragraph_language(
    paragraph,
    language: str,
    *,
    alignment: WD_ALIGN_PARAGRAPH | None = None,
) -> None:
    if alignment is not None:
        paragraph.alignment = alignment
    elif is_rtl(language):
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    if not is_rtl(language):
        return

    p_pr = paragraph._element.get_or_add_pPr()
    if p_pr.find(qn("w:bidi")) is None:
        p_pr.append(OxmlElement("w:bidi"))


def _style_run(
    run,
    language: str,
    *,
    bold: bool = False,
    size: int | None = None,
    color: RGBColor | None = None,
) -> None:
    run.font.bold = bold
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    _apply_run_font(run, language)


def _configure_document(doc: Document, language: str) -> None:
    font_name = _docx_font(language)
    for style_name in _DOCX_STYLES:
        try:
            style = doc.styles[style_name]
        except KeyError:
            continue
        style.font.name = font_name
        _set_rfonts(style.element.get_or_add_rPr(), font_name)


def _configure_page_margins(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin = _DOCX_PAGE_MARGIN
        section.bottom_margin = _DOCX_PAGE_MARGIN
        section.left_margin = _DOCX_PAGE_MARGIN
        section.right_margin = _DOCX_PAGE_MARGIN


def _cumulative_table_widths(n_years: int) -> list[float]:
    chart_area = _DOCX_CONTENT_WIDTH_IN - _DOCX_LABEL_COL_IN
    year_col = chart_area / n_years
    return [_DOCX_LABEL_COL_IN] + [year_col] * n_years


def _chart_area_width(n_years: int) -> float:
    return _DOCX_CONTENT_WIDTH_IN - _DOCX_LABEL_COL_IN


def _chart_render_width_px(n_years: int) -> int:
    return int(_chart_area_width(n_years) * _DOCX_CHART_DPI)


def _unavailable_row_widths() -> list[float]:
    return [_DOCX_LABEL_COL_IN, _DOCX_CONTENT_WIDTH_IN - _DOCX_LABEL_COL_IN]


def _donut_row_widths(*, has_target: bool) -> list[float]:
    if has_target:
        target_col = _DOCX_CONTENT_WIDTH_IN - _DOCX_LABEL_COL_IN - _DOCX_DONUT_COL_IN
        return [_DOCX_LABEL_COL_IN, _DOCX_DONUT_COL_IN, target_col]
    return [_DOCX_LABEL_COL_IN, _DOCX_DONUT_COL_IN]


def _donut_pair_widths() -> list[float]:
    label_col = (_DOCX_CONTENT_WIDTH_IN - (2 * _DOCX_DONUT_COL_IN)) / 2
    return [label_col, _DOCX_DONUT_COL_IN, label_col, _DOCX_DONUT_COL_IN]


def _set_cell_vertical_alignment(
    cell,
    alignment: WD_CELL_VERTICAL_ALIGNMENT = WD_CELL_VERTICAL_ALIGNMENT.CENTER,
) -> None:
    cell.vertical_alignment = alignment


def _center_table_cells(table) -> None:
    for row in table.rows:
        for cell in row.cells:
            _set_cell_vertical_alignment(cell)


def _style_heading_paragraph(paragraph, language: str, *, size: int, color: RGBColor | None = None) -> None:
    _apply_paragraph_language(paragraph, language)
    for run in paragraph.runs:
        _style_run(run, language, bold=True, size=size, color=color)


def _style_body_paragraph(
    paragraph,
    language: str,
    *,
    size: int | None = None,
    alignment: WD_ALIGN_PARAGRAPH | None = None,
) -> None:
    _apply_paragraph_language(paragraph, language, alignment=alignment)
    for run in paragraph.runs:
        _style_run(run, language, size=size)


def render_line_chart_asset(
    item: dict[str, Any],
    target_label: str,
    output_path: Path,
    *,
    width: int = CHART_WIDTH_PX,
    language: str = "English",
    show_labels: bool = True,
    session=None,
) -> Path:
    del session
    svg = render_line_chart_svg(
        item,
        width,
        chart_id="asset-line",
        show_value_labels=show_labels,
        show_target_labels=show_labels,
        target_label=target_label,
        language=language,
    )
    write_svg_png(svg, output_path, width=width, height=110)
    return output_path


def render_donut_asset(
    item: dict[str, Any],
    output_path: Path,
    *,
    language: str = "English",
    show_label: bool = True,
    session=None,
) -> Path:
    del session
    svg = render_donut_svg(item, show_label=show_label, language=language)
    write_svg_png(svg, output_path, width=64, height=64)
    return output_path


_TABLE_INNER_BORDER_COLOR = "C0C0C0"
_TABLE_INNER_BORDER_SIZE = 4  # eighths of a point (0.5 pt)


def _set_table_inner_borders(table) -> None:
    """Apply thin grey borders between cells only (no outer table border)."""
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    existing = tbl_pr.find(qn("w:tblBorders"))
    if existing is not None:
        tbl_pr.remove(existing)

    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "nil")
        borders.append(element)

    for edge in ("insideH", "insideV"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), str(_TABLE_INNER_BORDER_SIZE))
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), _TABLE_INNER_BORDER_COLOR)
        borders.append(element)

    tbl_pr.append(borders)


def _set_cell_text(
    cell,
    text: str,
    *,
    language: str = "English",
    bold: bool = False,
    size: int = _DOCX_TABLE_FONT,
    align_right: bool = False,
    align_center: bool = False,
) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    if align_center:
        alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align_right or is_rtl(language):
        alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        alignment = WD_ALIGN_PARAGRAPH.LEFT
    _apply_paragraph_language(p, language, alignment=alignment)
    run = p.add_run(text)
    _style_run(run, language, bold=bold, size=size)
    _set_cell_vertical_alignment(cell)


def _set_table_preferred_width(table, width_in: float) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(int(width_in * 1440)))
    tbl_w.set(qn("w:type"), "dxa")


def _set_column_widths(table, widths: list[float]) -> None:
    _set_table_preferred_width(table, sum(widths))
    for row in table.rows:
        for idx, width in enumerate(widths):
            if idx < len(row.cells):
                row.cells[idx].width = Inches(width)


def _add_data_row(
    table,
    row_idx: int,
    label: str,
    values: list[str],
    *,
    language: str = "English",
    bold_label: bool = False,
    bold_values: bool = False,
) -> None:
    _set_cell_text(table.cell(row_idx, 0), label, language=language, bold=bold_label, align_right=True)
    for col, value in enumerate(values, start=1):
        if col < len(table.rows[row_idx].cells):
            _set_cell_text(
                table.cell(row_idx, col),
                value,
                language=language,
                bold=bold_values,
                align_center=True,
            )


def _add_cumulative_block(
    doc: Document,
    item: dict[str, Any],
    labels: dict[str, str],
    target_label: str,
    assets_dir: Path,
    block_id: str,
    *,
    language: str = "English",
    session=None,
) -> None:
    if item.get("unavailable"):
        table = doc.add_table(rows=2, cols=2)
        _set_table_inner_borders(table)
        _set_table_fixed_layout(table)
        _set_column_widths(table, _unavailable_row_widths())
        _set_cell_text(table.cell(0, 0), item["label"], language=language)
        merged = table.cell(0, 1)
        merged.merge(table.cell(1, 1))
        _set_cell_text(
            merged,
            item.get("unavailable_label") or not_available(language),
            language=language,
            align_center=True,
        )
        table.cell(1, 0).text = ""
        _set_cell_vertical_alignment(table.cell(1, 0))
        _center_table_cells(table)
        doc.add_paragraph("")
        return

    n_years = len(item["years"])
    chart_path = assets_dir / f"{block_id}_line.png"
    chart_width_px = _chart_render_width_px(n_years)
    render_line_chart_asset(
        item,
        target_label,
        chart_path,
        width=chart_width_px,
        language=language,
        session=session,
    )

    show_reporting, show_implementing = cumulative_table_rows(item)
    n_rows = 2 + int(show_reporting) + int(show_implementing)
    table = doc.add_table(rows=n_rows, cols=n_years + 1)
    _set_table_inner_borders(table)
    _set_table_fixed_layout(table)
    _set_column_widths(table, _cumulative_table_widths(n_years))

    _set_cell_text(table.cell(0, 0), item["label"], language=language)
    chart_cell = table.cell(0, 1)
    chart_cell.merge(table.cell(0, n_years))
    chart_cell.text = ""
    _apply_paragraph_language(chart_cell.paragraphs[0], language, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    chart_cell.paragraphs[0].add_run().add_picture(
        str(chart_path),
        width=Inches(_chart_area_width(n_years)),
    )
    _set_cell_vertical_alignment(chart_cell)

    _add_data_row(table, 1, labels["year"], item["years"], language=language, bold_label=True, bold_values=True)
    row_idx = 2
    if show_reporting:
        _add_data_row(table, row_idx, labels["reporting"], item["reporting"], language=language)
        row_idx += 1
    if show_implementing:
        _add_data_row(table, row_idx, labels["implementing"], item["implementing"], language=language)

    _center_table_cells(table)
    doc.add_paragraph("")


def _set_table_fixed_layout(table) -> None:
    """Use fixed column widths so Word does not reflow cells onto separate lines."""
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")


def _set_row_cant_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def _add_donut_image_cell(cell, image_path: Path, *, language: str) -> None:
    cell.text = ""
    _apply_paragraph_language(cell.paragraphs[0], language, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    cell.paragraphs[0].add_run().add_picture(str(image_path), width=Inches(0.45))
    _set_cell_vertical_alignment(cell)


def _add_donut_block(
    doc: Document,
    item: dict[str, Any],
    assets_dir: Path,
    block_id: str,
    *,
    language: str = "English",
    session=None,
) -> None:
    if item.get("unavailable"):
        table = doc.add_table(rows=1, cols=2)
        _set_table_inner_borders(table)
        _set_table_fixed_layout(table)
        _set_column_widths(table, _unavailable_row_widths())
        _set_cell_text(table.cell(0, 0), item["label"], language=language)
        _set_cell_text(
            table.cell(0, 1),
            item.get("unavailable_label") or not_available(language),
            language=language,
            align_center=True,
        )
        _center_table_cells(table)
        doc.add_paragraph("")
        return

    donut_path = assets_dir / f"{block_id}_donut.png"
    render_donut_asset(item, donut_path, language=language, session=session)

    has_target = bool(item.get("target_label"))
    table = doc.add_table(rows=1, cols=3 if has_target else 2)
    _set_table_inner_borders(table)
    _set_table_fixed_layout(table)
    _set_column_widths(table, _donut_row_widths(has_target=has_target))

    _set_cell_text(table.cell(0, 0), item["label"], language=language)
    _add_donut_image_cell(table.cell(0, 1), donut_path, language=language)
    if has_target:
        _set_cell_text(
            table.cell(0, 2),
            item["target_label"].replace("\n", " "),
            language=language,
        )

    _center_table_cells(table)
    doc.add_paragraph("")


def _add_donut_pair_block(
    doc: Document,
    items: list[dict[str, Any]],
    assets_dir: Path,
    block_id: str,
    *,
    language: str = "English",
    session=None,
) -> None:
    left, right = items[0], items[1]
    if left.get("unavailable") and right.get("unavailable"):
        for idx, item in enumerate((left, right)):
            _add_donut_block(doc, item, assets_dir, f"{block_id}_{idx}", language=language, session=session)
        return

    left_path = assets_dir / f"{block_id}_left_donut.png"
    right_path = assets_dir / f"{block_id}_right_donut.png"
    if not left.get("unavailable"):
        render_donut_asset(left, left_path, language=language, session=session)
    if not right.get("unavailable"):
        render_donut_asset(right, right_path, language=language, session=session)

    table = doc.add_table(rows=1, cols=4)
    _set_table_inner_borders(table)
    _set_table_fixed_layout(table)
    _set_column_widths(table, _donut_pair_widths())
    _set_row_cant_split(table.rows[0])

    _set_cell_text(table.cell(0, 0), left["label"], language=language)
    if left.get("unavailable"):
        _set_cell_text(
            table.cell(0, 1),
            left.get("unavailable_label") or not_available(language),
            language=language,
            align_center=True,
        )
    else:
        _add_donut_image_cell(table.cell(0, 1), left_path, language=language)
    _set_cell_text(table.cell(0, 2), right["label"], language=language)
    if right.get("unavailable"):
        _set_cell_text(
            table.cell(0, 3),
            right.get("unavailable_label") or not_available(language),
            language=language,
            align_center=True,
        )
    else:
        _add_donut_image_cell(table.cell(0, 3), right_path, language=language)

    _center_table_cells(table)
    doc.add_paragraph("")


def _add_sp_section(
    doc: Document,
    payload: dict[str, Any],
    assets_dir: Path,
    *,
    show_section_title: bool = True,
) -> None:
    language = payload.get("language", "English")
    if show_section_title:
        title = doc.add_paragraph(payload["title"])
        _style_heading_paragraph(title, language, size=12, color=IFRC_RED)

    labels = payload["table_labels"]
    target_label = payload["headers"]["target"]

    for idx, item in enumerate(payload["cumulative"]):
        _add_cumulative_block(
            doc, item, labels, target_label, assets_dir, f'{payload["section"]}_cum_{idx}',
            language=language,
        )

    for row_idx, pair in enumerate(payload.get("donut_pairs", [])):
        if len(pair) >= 2:
            _add_donut_pair_block(
                doc, pair[:2], assets_dir, f'{payload["section"]}_pair_{row_idx}',
                language=language,
            )
        elif pair:
            _add_donut_block(
                doc, pair[0], assets_dir, f'{payload["section"]}_pair_{row_idx}_0',
                language=language,
            )

    foot = doc.add_paragraph(payload["footnote"])
    _style_body_paragraph(foot, language, size=8)


def render_report_docx(
    model,
    *,
    language: str = "English",
    output_path: Path,
    sections: list[str] | None = None,
    mapping=None,
) -> Path:
    titles = report_titles()
    report_title = titles.get(language, titles["English"])
    parts = report_parts()
    section_order = sections or [s for part in parts for s in part["sections"]]
    full_mapping = mapping if mapping is not None else mapping_from_model(model)

    doc = Document()
    _configure_document(doc, language)
    _configure_page_margins(doc)
    heading = doc.add_heading(report_title, level=0)
    _style_heading_paragraph(heading, language, size=16, color=IFRC_RED)
    doc.add_paragraph("")

    with tempfile.TemporaryDirectory(prefix="pb_assets_") as tmp:
        assets_dir = Path(tmp)
        for part in parts:
            part_title = part["title"].get(language, part["title"]["English"])
            part_heading = doc.add_heading(part_title, level=1)
            _style_heading_paragraph(part_heading, language, size=14, color=IFRC_RED)

            for section in part["sections"]:
                if section not in section_order:
                    continue
                if not section_has_indicators(full_mapping, section):
                    continue
                payload = build_payload(model, section, language, mapping=full_mapping)
                _add_sp_section(
                    doc,
                    payload,
                    assets_dir,
                    show_section_title=not section_uses_part_heading_only(part["id"]),
                )

    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path
